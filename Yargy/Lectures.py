
import yargy
import os
import re
import copy
import glob
import json
from docx import Document
from yargy import Parser, rule, and_, or_
from yargy.pipelines import morph_pipeline
from yargy.predicates import gram, is_capitalized, dictionary, is_upper, length_eq


class LectureParser:

    def __init__(self,name=None):

        self.section_rule = rule(
            dictionary(
                {
                    "раздел", "тема", "дисциплина", "наименование"
                }))

        self.lectures_rule = rule(
            morph_pipeline([
                'тема лекций',
                'содержание занятий',
                'содержание лекционного занятия'
            ]))

        self.pract_rule = rule(
            morph_pipeline(
                [
                'наименование'
                ]))


        self.srs_rule = rule(
             morph_pipeline(
                [
                'СРС'
                ]))


        self.docxdoc=DocumentPrepare(name).open_doc()


    def sections(self, segment='all'):
        """
        attributes:
        <all>
        <themes>
        <lectures>
        <practices>
        <srs>
        """

        themes = Parser(self.section_rule)
        lectures = Parser(self.lectures_rule)
        practices = Parser(self.pract_rule)
        srs = Parser(self.srs_rule)
        found = False
        for table in self.docxdoc.tables:
            for column in table.columns:
                for cell in column.cells:
                    index = 0

                    if segment=='all' or segment=='themes':
                        cell_search_themes = themes.findall(cell.text)
                        for each in cell_search_themes:
                            index+=1
                        if index > 2:
                            return self.lectures(table,column)
                            if segment !='all':
                                found = True
                            print("this is theme")
                            break 

                    if segment=='all' or segment=='lectures':
                        cell_search_lectures = lectures.findall(cell.text)
                        for each in cell_search_lectures:
                            return self.lectures(table,column)
                            if segment !='all':
                                found = True
                            print("ЛЕКЦИИ")
                            break

                    if segment=='all' or segment=='practices':
                        cell_search_practices = practices.findall(cell.text)
                        for each in cell_search_practices:
                            return self.lectures(table,column)
                            if segment !='all':
                                found = True
                            print("практика")
                            break

                    if segment=='all' or segment=='srs':
                        cell_search_srs = srs.findall(cell.text)
                        for each in cell_search_srs:
                            return self.lectures(table,column)
                            if segment !='all':
                                found = True
                            print("практика")
                            break
                        
                    
                if found: break
            if found: break


    def lectures(self,table, column):
        key = column.cells[0].text
        lect_dict={key:[]}
        flag = False
        separator =False
        save_pre=''
        save_lect=''
        for cell in column.cells:
            lect = cell.text
            
            for row in table.rows:
                for cell in row.cells:

                    if (cell.text == "Контроль") or (cell.text == "Всего:") or (cell.text == "Итого") :
                        break

                    if flag:
                    
                        precision = cell.text
                        if save_lect != lecture:
                            lect_dict[key].append(lecture+'=')
                            save_lect = lecture

                        if re.sub(r'[^\w\s]+|[\d]+',r'',precision).strip() != '' and save_pre!=precision:
                            lect_dict[key].append(precision+'|')
                            save_pre = precision
                            
                        flag = False

                    if (cell.text == lect) and lect != '':
                        lecture = cell.text
                        flag = True

        return(lect_dict)

class DocumentPrepare:

    def __init__(self,name='None'):
        self.namedoc=name

    def open_doc(self, docname='None'):
        
        if self.namedoc is not 'None':
            docname = self.namedoc
        if (docname is 'None') and (self.namedoc is 'None'):
            print("print the name of the file (yargy)")
            docname = input()
        try: 
            docx = Document(docname)

        except PackageNotFoundError:
            if os.name == 'nt':
                path = "docs\\" + name 
            if os.name == 'posix':
                path = os.getcwd() + "/Git/parser_results_and_competitions/co-co-corpus" +docname #  Linux
            docx = Document(path)

        return docx


#mytext = LectureParser('/ЧелГУ/5_РПД _Математический анализ, Дифференциальные и разностные уравнения.docx')
#mytext = LectureParser('/ЮГРА/15.Сети ЭВМ и телекоммуникации.docx')
#mytext = LectureParser('/ЮУрГУ/РПД Архитектура ЭВМ (09.03.01, 2016, (4.0), Информатика и вычислительная техника(19610)).docx')
#mytext = LectureParser('/УрФУ/5_Раб.программа дисциплины Инструм моделирования БП.docx')
#mytext.sections('lectures')

def test_search(univer='all',code='all'):
    unilist = ['ЧелГУ','ЮУрГУ','ЮГРА','УрФУ']
    if univer != 'all':
        unilist = [univer]

    for uni in unilist:
        print(uni)
        path = os.getcwd() + "/Git/parser_results_and_competitions/co-co-corpus/"+uni+"/"
        for file in glob.glob(os.path.join(path, '*.docx')):
            mytext = LectureParser(file)
            savetojson = mytext.sections(code)
            with open('data.txt','w') as outfile:
                json.dump(savetojson, outfile)
            print('end of file')

test_search()
