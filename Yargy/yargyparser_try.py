from docx import Document
from docx.shared import Inches
import yargy
import os
from yargy.tokenizer import Tokenizer, TokenRule
from yargy.pipelines import morph_pipeline
import input_test as inp

class CompetitionResult():

    def __init__(self, FGOS, competition, result=None):
        self.FGOS = FGOS
        self.competition = competition
        self.result = result




class YargyParser:

    def __init__(self, filename="none"):
        if filename == "none":
            print("print the name of the file (yargy)")
            filename = input()
        self.filename = filename
        self.document = Document(os.getcwd() + "\\" + (self.filename + ".docx"))


    def text_prepare(self):
        my_table = ""
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    my_table += cell.text
                    my_table += "~"
                my_table += "@"

        text = my_table
        return text


    def search_place_FGOS(self):

        FGOS_list = self.token_FGOS(self.text_prepare())
        entitis = []
        if len(FGOS_list)>0:
            k = FGOS_list[0]
            print(len(FGOS_list))
            for i in FGOS_list:
                text = self.text_prepare()
                # print(i.value)
                k = i
                if i.value not in entitis:
                    pr = text[k.span[1]:text.find('@',k.span[1])]
                    entitis.append(i.value)
                    print(i.value)
                    #print(pr)
                    competence, know, can, own = self.separate(pr)
                    print(competence)
        else:
            checklist = 0
            for i in self.document.paragraphs:

                text = i.text
                FGOS_list += self.token_FGOS(text)
                len_list = len(FGOS_list)
                if len_list != checklist:
                    print(i.text)
                    checklist = len_list
                '''
                надо сохранять строку в которой найдет код
                сохранять ее как компетенции
                '''
            print(len(FGOS_list))






    def separate(self, text):
        competence = text[1:text.find('~',1)]
        know = ""
        can = ""
        own = ""
        return competence, know,can,own



    def token_FGOS(self, text):
        tokenizer = Tokenizer()
        FOS_RULE = TokenRule('FOS', '[А-Я]+К+-+[0-9]+') #  букв не больше 3 и последняя к
        tokenizer.remove_types('EOL', 'RU','PUNCT','OTHER','INT','LATIN')
        tokenizer.add_rules(FOS_RULE)
        return list(tokenizer(text))

'''
пробовали на:

РПД_ПрИнж_Теория автоматов
ИС в предметной области+
15.Сети ЭВМ и телекоммуникации
25_РПД Разработка приложений для работы с БД

'''



mytry = YargyParser("РПД_ПрИнж_Теория автоматов")
mytry.search_place_FGOS()
#  mytry.token_FOS()
