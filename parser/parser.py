from yargy import Parser, rule, and_, or_
from yargy.predicates import gram, is_capitalized, dictionary, is_upper, length_eq
from docx import Document
from docx.shared import Inches
import list_header as lh

list_header = lh.list_header

class CompetitionResult():
    def __init__(self, FGOS, competition, result):
        self.FGOS = FGOS
        self.competition = competition
        self.result = result

FGOS = rule(
     and_(
         is_upper(),
         or_(
            length_eq(2),
            length_eq(3)
         )
        )
)

IsCodeFGOS = rule(
    dictionary(
    {'код', 'ФГОС'})
)
IsCompetitions = rule(dictionary(
    {
    'компетенция',
    'содержание'
    }
))
IsResults = rule(dictionary(
    {
        'результат'
    }
))
TableIsCompetitionsAndResults = rule(dictionary(
    {
        'результат',
        'ФГОС',
        'компетенция'
    }
))

parser_Table = Parser(TableIsCompetitionsAndResults)
parser_FGOS = Parser(IsCodeFGOS)
parser_Result = Parser(IsResults)
parser_Competition = Parser(IsCompetitions)
parser_code_FGOS = Parser(FGOS)

# из текста
def FindFGOSs(text):
    list_code_fgos = []
    for match in parser_code_FGOS.findall(text):
        index = 0
        for i in match.tokens:
            code_FGOS = match.tokens[index].value+text[match.tokens[index].span[1]]+text[match.tokens[index].span[1]+1]
            if text[match.tokens[index].span[1]+2].isdigit():
                code_FGOS = code_FGOS + text[match.tokens[index].span[1]+2]
            list_code_fgos.append(code_FGOS)
            index +=1
    return list_code_fgos

#из строчки
def FindFGOS(text):
    code_FGOS = ""
    for match in parser_code_FGOS.findall(text):
        code_FGOS = match.tokens[0].value+text[match.tokens[0].span[1]]+text[match.tokens[0].span[1]+1]
        if text[match.tokens[0].span[1]+2].isdigit():
            code_FGOS = code_FGOS + text[match.tokens[0].span[1]+2]
    return code_FGOS


document = Document('C://Users//Хитрый//Documents//GitHub//corpus//Корпус документов//ЧелГУ ИИТ//2016//РПД ПИ 2014 год набора//РПД Государственная итоговая аттестация.docx')
#document = Document('C://Users//Хитрый//Documents//GitHub//parser_results_and_competitions//parser//25_РПД Разработка приложений для работы с БД.docx')
#document = Document('C://Users//Хитрый//Documents//GitHub//parser_results_and_competitions//parser//ОП ВО 09.03.01, 2016, (4.0), Информатика и вычислительная техника (19610).docx')


def FindTable():
    for table in document.tables:

        data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            data.append(row_data)


        for j in data[0]:
            if parser_Table.find(j) is not None:
                return data

FGOS = []
competitions = []
results = []

table = FindTable()
if table is not None:
    #определение какой столбец за что отвечает
    key = 0
    dict={}
    for i in table[0]:
        print (i)
        if parser_FGOS.find(i) is not None:
            dict['ФГОС'] = i
        elif parser_Competition.find(i) is not None:
            dict['компетенции'] = i
        elif parser_Result.find(i) is not None:
            dict['результаты'] = i
    print(dict)


    for i in table:#i - словарь

        if 'ФГОС' in dict:
            FGOS.append(i[dict['ФГОС']])

        if 'компетенции' in dict and not i[dict['компетенции']] == dict['компетенции']:
            competitions.append(i[dict['компетенции']])

        if 'результаты' in dict:
            results.append(i[dict['результаты']])


    if len(results) != 0:
        if len(competitions)!= 0 and len(FGOS) != 0:
            competition_result = []
            index_list = 0
            for i in competitions:
                competition_result.append(CompetitionResult(FGOS[index_list],competitions[index_list], results[index_list]))
                index_list += 1
            for i in competition_result:
                print('Код: ' + i.FGOS, 'Компетенция: ' + i.competition, 'Результат: ' + i.result)
        elif len(competitions) == 0 and len(FGOS) != 0:
            print('если коды и компетенции в столбике коды')

        elif len(competitions) != 0 and len(FGOS) == 0:
            print('если коды и компетенции в столбике компетенции')

            competition_result = []
            index_list = 0
            for i in competitions:
                code = FindFGOS(i)
                index = i.find(code)
                competition_result.append(CompetitionResult(code, i[index+len(code):], results[index_list]))
                index_list += 1

            for i in competition_result:
                print('Код: ' + i.FGOS, 'Компетенция: '+ i.competition, 'Результат: '+i.result)

    print(FGOS)
    print(competitions)
    print(results)

else:
    text = ""
    index = 0
    for i in document.paragraphs:
        if document.paragraphs[index].text in list_header:
            print(document.paragraphs[index].text)
            if parser_Table.find(document.paragraphs[index].text ) is not None:
                print(parser_Table.find(document.paragraphs[index].text))
                while document.paragraphs[index+1].text not in list_header:
                    print(document.paragraphs[index+1].text)
                    print(document.paragraphs[index+1].text not in list_header)
                    text += document.paragraphs[index+1].text
                    index +=1
        index += 1
    print(text)




